from django.shortcuts import redirect, render
from django.contrib import messages
from django.views import View
import pandas as pd
import django_tables2 as tables
from django_filters.views import FilterView
from django_tables2.views import SingleTableMixin, SingleTableView
from app.tables import BillTable, FileTable, ParentTable, PlpRecordTable, PlpTable, StudentTable, SubjectDocTable, SubjectTable, UserTable
from app.models import bills, parents, plprecord, plps, students, subject_doc, subjects, uploadrecords, user
from app.filters import BillFilter, ParentFilter, PlpFilter, PlpRecordFilter, StudentFilter, SubjectDocFilter, UserFilter
from django.views.generic import CreateView, DeleteView, UpdateView,FormView
#Raw Anvarisy
from django.views.generic.edit import BaseDeleteView
from django.views.generic.list import MultipleObjectTemplateResponseMixin
#End
from django.urls import reverse_lazy
from django.contrib.auth.models import Permission
from django.contrib.auth.mixins import LoginRequiredMixin,PermissionRequiredMixin
from django.contrib.auth.views import LoginView, LogoutView
from django.contrib.auth import logout
from django.contrib.auth.forms import UserCreationForm,PasswordChangeForm
from app.forms import BillForm, FileForm, ParentForm, PlpForm, PlpRecordForm, RaporForm, RegisterForm, StudentForm, SubjectDocForm, SubjectForm, UpdateForm
from django.contrib.auth.decorators import permission_required  
from django.utils.decorators import method_decorator
from django.contrib.auth.forms import PasswordChangeForm
from django.contrib.auth.views import PasswordChangeView
from django.http import HttpResponse, HttpResponseRedirect
from django.contrib import messages
from mailmerge import MailMerge
import os
from school_system.settings import  BASE_DIR
import json
from datetime import datetime
from docx import Document
from docxcompose.composer import Composer
# from pyexcel_xls import get_data as xls_get
# from pyexcel_xlsx import get_data as xlsx_get
# Create your views here.
class ViewHomepage(View):
    template_name='pages/homepage.html'
    def get(self, request):
        return render(request, self.template_name,{'title':'Home'})

#-----------------------User Manager--------------------
class ListUser(PermissionRequiredMixin,LoginRequiredMixin,SingleTableMixin, FilterView):
    permission_required = ('app.im_admin')
    login_url = '/login/'
    table_class = UserTable
    queryset = user.objects.all()
    template_name = "pages/index.html"
    filterset_class = UserFilter

class AddUser(PermissionRequiredMixin,LoginRequiredMixin, CreateView):
    permission_required = ('app.im_admin')
    form_class = RegisterForm
    login_url = '/login/' 
    template_name = 'pages/form.html'
    success_url = reverse_lazy('list-user')
    raise_exception = True
    def get_context_data(self, **kwargs):
        context = super(AddUser, self).get_context_data(**kwargs)
        context['title'] = 'Add User'
        context['button'] = 'Add User'
        return context
    def form_valid(self, form):
        u = form.save()
        perm = Permission.objects.get(
        codename='im_admin')
        print(perm)
        us = user.objects.get(email=u)
        if us.is_admin:
            us.user_permissions.add(perm)
        return redirect('list-user')

class UpdateUser(UpdateView, LoginRequiredMixin,PermissionRequiredMixin):
    permission_required = ('app.im_admin')
    template_name = 'pages/form.html'
    success_url = reverse_lazy('list-user')
    form_class = UpdateForm
    model = user
    login_url = '/login/'
    def get_context_data(self, **kwargs):
        context = super(UpdateUser, self).get_context_data(**kwargs)
        context['title'] = 'Update User'
        context['button'] = 'Update User'
        return context
    def form_valid(self, form):
        u = form.save()
        perm = Permission.objects.get(
        codename='im_admin')
        print(perm)
        us = user.objects.get(email=u)
        if us.is_admin:
            us.user_permissions.add(perm)
        else:
            try:
                us.user_permissions.remove(perm)
            except:
                pass
        return redirect('list-user')

class DeleteUser(PermissionRequiredMixin, DeleteView):
    permission_required = ('app.im_admin')
    raise_exception = True
    login_url = '/login/' 
    model = user
    success_url = reverse_lazy('list-user')
        
class UpdatePassword(PasswordChangeView, LoginRequiredMixin):
    success_url = reverse_lazy('home')
    form_class = PasswordChangeForm
    login_url = '/login/'

class ViewLogin(LoginView):
    template_name = 'pages/login_page.html'
    success_url = reverse_lazy('home')

class ViewLogout(LogoutView):
    success_url = reverse_lazy('login')
    template_name = 'app/user_confirm_logout.html'

#--------------------------End of User--------------- 


#---------------------Parent-----------------------
class ListParent(LoginRequiredMixin,SingleTableMixin, FilterView):
    login_url = '/login/'
    table_class = ParentTable
    queryset = parents.objects.all()
    template_name = "pages/index.html"
    filterset_class = ParentFilter
    def get_context_data(self, **kwargs):
        context = super(ListParent, self).get_context_data(**kwargs)
        context['del'] = '/del-par/'
        context['imported'] = True
        context['post_import']='/import-parent/'
        return context

class AddParent(LoginRequiredMixin, CreateView):
    form_class = ParentForm
    login_url = '/login/' 
    template_name = 'pages/form.html'
    success_url = reverse_lazy('list-parent')
    def get_context_data(self, **kwargs):
        context = super(AddParent, self).get_context_data(**kwargs)
        context['title'] = 'Add Parent'
        context['button'] = 'Add Parent'
        return context

class UpdateParent(UpdateView, LoginRequiredMixin):
    template_name = 'pages/form.html'
    success_url = reverse_lazy('list-parent')
    form_class = ParentForm
    model = parents
    login_url = '/login/'
    def get_context_data(self, **kwargs):
        context = super(UpdateParent, self).get_context_data(**kwargs)
        context['title'] = 'Update Parent'
        context['button'] = 'Update Parent'
        return context
 
class DeleteParent(LoginRequiredMixin, DeleteView):
    login_url = '/login/' 
    model = parents
    template_name = "app/confirm_delete.html"
    success_url = reverse_lazy('list-parent')

class DeleteAllParent(LoginRequiredMixin, View):
    # models = parents
    # success_url = reverse_lazy('list-parent')
    # template_name_suffix = '_confirm_delete'
    # queryset = parents.objects.all()
    def post(self, request):
        par_arr = request.POST.getlist('selection')
        for email in par_arr:
            parents.objects.filter(parent_email=email).delete()
        return redirect('list-parent') 

class ImportParent(LoginRequiredMixin, View):
    def post(self, request):
        try:
            data = pd.read_excel(request.FILES.get('file_import'))
            df = pd.DataFrame(data)
            # _parents = data[['parent_name','parent_mobile','parent_email','parent_add']]
            data_json =[]
            df_json = df.to_json(orient='records',force_ascii=False)
            d = json.loads(df_json)
            for item in d:
                try:
                    o = parents.objects.create(**item)
                    o.save()
                except Exception as e:
                    print(e)
            # for item in data['parent_name']:
            #     print(item)
            # data = None
            # file = request.FILES['file_import']
            # if (str(file).split('.')[-1] == 'xls'):
            #     data = xls_get(file, column_limit=5)
            # elif (str(file).split('.')[-1] == 'xlsx'):
            #     data = xlsx_get(file, column_limit=5)
            # else:
            #     messages.add_message(request,messages.ERROR,'File harus berupa xls atau xlsx')
            #     return redirect('list-parent')
            # _parents = data['Parent']
            # for parent in _parents:
            #     print(parent)
            # # for item in data:
            # #     print(data)
            return redirect('list-parent')
        except Exception as e:
            messages.add_message(request,messages.ERROR,'File tidak terbaca, pastikan nama sheet adalah Parent')
            return redirect('list-parent')
        # return redirect('list-parent')
#-----------------------End Of Parent-----------------


#--------------------File (Belum selesai)-------------
class ListFile(LoginRequiredMixin, SingleTableView):
    login_url = '/login/'
    table_class = FileTable
    queryset = uploadrecords.objects.all()
    template_name = "pages/index.html"

class AddFile(LoginRequiredMixin,CreateView):
    form_class = FileForm
    login_url = '/login/' 
    template_name = 'pages/form.html'
    success_url = reverse_lazy('list-file')
    raise_exception = True
    def get_context_data(self, **kwargs):
        context = super(AddFile, self).get_context_data(**kwargs)
        context['title'] = 'Add File to be Imported'
        context['button'] = 'Upload'
        return context
    def form_valid(self, form):
        f = form.save(commit=False)
        f.uploaded_by = self.request.user
        f.save()
        return redirect('list-file')

#------------------------Student-----------------------
class ListStudent(LoginRequiredMixin,SingleTableMixin, FilterView):
    login_url = '/login/'
    table_class = StudentTable
    queryset = students.objects.all()
    template_name = "pages/index.html"
    filterset_class = StudentFilter
    def get_context_data(self, **kwargs):
        context = super(ListStudent, self).get_context_data(**kwargs)
        context['del'] = '/del-stud/'
        context['imported'] = True
        context['post_import']='/import-student/'
        return context

class DeleteAllStudent(View):
    # models = parents
    # success_url = reverse_lazy('list-parent')
    # template_name_suffix = '_confirm_delete'
    # queryset = parents.objects.all()
    def post(self, request):
        stud_arr = request.POST.getlist('selection')
        for nis in stud_arr:
            students.objects.filter(nis_student=nis).delete()
        return redirect('list-student') 

class AddStudent(LoginRequiredMixin, CreateView):
    form_class = StudentForm
    login_url = '/login/' 
    template_name = 'pages/form.html'
    success_url = reverse_lazy('list-student')
    success_message = 'Doc successfully created!'
    error_message = "Error saving the Doc, check fields below."
    def get_context_data(self, **kwargs):
        context = super(AddStudent, self).get_context_data(**kwargs)
        context['title'] = 'Add Student'
        context['button'] = 'Add Student'
        return context

class UpdateStudent(UpdateView, LoginRequiredMixin):
    template_name = 'pages/form.html'
    success_url = reverse_lazy('list-student')
    form_class = StudentForm
    model = students
    login_url = '/login/'
    def get_context_data(self, **kwargs):
        context = super(UpdateStudent, self).get_context_data(**kwargs)
        context['title'] = 'Update Student'
        context['button'] = 'Update Student'
        return context
 
class DeleteStudent(LoginRequiredMixin, DeleteView):
    login_url = '/login/' 
    model = students
    template_name = "app/confirm_delete.html"
    success_url = reverse_lazy('list-student')

class ImportStudent(LoginRequiredMixin, View):
    def post(self, request):
        try:
            data = pd.read_excel(request.FILES.get('file_import'))
            
            df = pd.DataFrame(data)
            # pd.to_datetime(df['dob_student']).apply(lambda x: x.date())
            df = df.applymap(str)
            df.replace("nan","",inplace=True)
            # df['json'] = df.apply(lambda x: x.to_json(), axis=1)
            df_json = df.to_json(orient='records',force_ascii=False)
            # print(df_json)
            a = json.loads(df_json)
            for item in a:
                try:
                    _parent = item.pop('parent')
                    a = item.pop('dob_student')
                    _dob = datetime.strptime(str(a), '%Y-%m-%d %H:%M:%S')
                    # dob = _dob.strftime('%Y-%m-%d')
                    # print(dob)
                    # print(datetime.strptime(item['dob_student'], '%Y-%m-%d'))
                    o = students.objects.create(**item)
                    o.parent_id = _parent
                    o.dob_student = _dob
                    # # o.dob_student
                    o.save()
                except Exception as e:
                    print(e)
            return redirect('list-student')
        except Exception as e:
            messages.add_message(request,messages.ERROR,'File tidak terbaca, pastikan nama sheet adalah Parent')
            return redirect('list-student')
#-----------------BILL------------------------
class ListBill(LoginRequiredMixin,SingleTableMixin, FilterView):
    login_url = '/login/'
    table_class = BillTable
    model = bills
    queryset = bills.objects.all()
    template_name = "pages/index.html"
    filterset_class = BillFilter
    def get_context_data(self, **kwargs):
        context = super(ListBill, self).get_context_data(**kwargs)
        context['del'] = '/del-bill/'
        return context

class DeleteAllBill(View):
    # models = parents
    # success_url = reverse_lazy('list-parent')
    # template_name_suffix = '_confirm_delete'
    # queryset = parents.objects.all()
    def post(self, request):
        bill_arr = request.POST.getlist('selection')
        for _id in bill_arr:
            bills.objects.filter(id=_id).delete()
        return redirect('list-bill') 

class AddBill(LoginRequiredMixin, CreateView):
    form_class = BillForm
    login_url = '/login/' 
    template_name = 'pages/form.html'
    success_url = reverse_lazy('list-bill')
    success_message = 'Doc successfully created!'
    error_message = "Error saving the Doc, check fields below."
    def get_context_data(self, **kwargs):
        context = super(AddBill, self).get_context_data(**kwargs)
        context['title'] = 'Add Bill'
        context['button'] = 'Add Bill'
        return context

class UpdateBill(UpdateView, LoginRequiredMixin):
    template_name = 'pages/form.html'
    success_url = reverse_lazy('list-bill')
    form_class = BillForm
    model = bills
    login_url = '/login/'
    def get_context_data(self, **kwargs):
        context = super(UpdateBill, self).get_context_data(**kwargs)
        context['title'] = 'Update Bill'
        context['button'] = 'Update Bill'
        return context
 
class DeleteBill(LoginRequiredMixin, DeleteView):
    login_url = '/login/' 
    model = bills
    template_name = "app/confirm_delete.html"
    success_url = reverse_lazy('list-bill')
#-----------------------End of Bill--------------------

#------------------------PLP--------------------------
class ListPlp(LoginRequiredMixin,SingleTableMixin, FilterView):
    login_url = '/login/'
    table_class = PlpTable
    model = plps
    queryset = plps.objects.all()
    template_name = "pages/index.html"
    filterset_class = PlpFilter
    def get_context_data(self, **kwargs):
        context = super(ListPlp, self).get_context_data(**kwargs)
        context['del'] = '/del-plp/'
        return context

class DeleteAllPlp(View):
    # models = parents
    # success_url = reverse_lazy('list-parent')
    # template_name_suffix = '_confirm_delete'
    # queryset = parents.objects.all()
    def post(self, request):
        plp_arr = request.POST.getlist('selection')
        for _id in plp_arr:
            plps.objects.filter(plp_code=_id).delete()
        return redirect('list-plp') 

class AddPlp(LoginRequiredMixin, CreateView):
    form_class = PlpForm
    login_url = '/login/' 
    template_name = 'pages/form.html'
    success_url = reverse_lazy('list-plp')
    success_message = 'Doc successfully created!'
    error_message = "Error saving the Doc, check fields below."
    def get_context_data(self, **kwargs):
        context = super(AddPlp, self).get_context_data(**kwargs)
        context['title'] = 'Add PLP'
        context['button'] = 'Add PLP'
        return context

class UpdatePlp(UpdateView, LoginRequiredMixin):
    template_name = 'pages/form.html'
    success_url = reverse_lazy('list-plp')
    form_class = PlpForm
    model = plps
    login_url = '/login/'
    def get_context_data(self, **kwargs):
        context = super(UpdatePlp, self).get_context_data(**kwargs)
        context['title'] = 'Update Plp'
        context['button'] = 'Update Plp'
        return context
 
class DeletePlp(LoginRequiredMixin, DeleteView):
    login_url = '/login/' 
    model = plps
    template_name = "app/confirm_delete.html"
    success_url = reverse_lazy('list-plp')

#----------------------------Record PLP--------------------
class ListPlpRecord(LoginRequiredMixin,SingleTableMixin, FilterView):
    login_url = '/login/'
    table_class = PlpRecordTable
    model = plprecord
    queryset = plprecord.objects.all()
    template_name = "pages/index.html"
    filterset_class = PlpRecordFilter
    def get_context_data(self, **kwargs):
        context = super(ListPlpRecord, self).get_context_data(**kwargs)
        context['del'] = '/del-plp-record/'
        return context

class DeleteAllPlpRecord(View):
    # models = parents
    # success_url = reverse_lazy('list-parent')
    # template_name_suffix = '_confirm_delete'
    # queryset = parents.objects.all()
    def post(self, request):
        _arr = request.POST.getlist('selection')
        for _id in _arr:
            plprecord.objects.filter(id=_id).delete()
        return redirect('list-plp-record') 

class AddPlpRecord(LoginRequiredMixin, CreateView):
    form_class = PlpRecordForm
    login_url = '/login/' 
    template_name = 'pages/form.html'
    success_url = reverse_lazy('list-plp-record')
    success_message = 'Doc successfully created!'
    error_message = "Error saving the Doc, check fields below."
    def get_context_data(self, **kwargs):
        context = super(AddPlpRecord, self).get_context_data(**kwargs)
        context['title'] = 'Add Record'
        context['button'] = 'Add Record'
        return context

class UpdatePlpRecord(UpdateView, LoginRequiredMixin):
    template_name = 'pages/form.html'
    success_url = reverse_lazy('list-plp-record')
    form_class = PlpRecordForm
    model = plprecord
    login_url = '/login/'
    def get_context_data(self, **kwargs):
        context = super(UpdatePlpRecord, self).get_context_data(**kwargs)
        context['title'] = 'Update Record'
        context['button'] = 'Update Record'
        return context
 
class DeletePlpRecord(LoginRequiredMixin, DeleteView):
    login_url = '/login/' 
    model = plprecord
    template_name = "app/confirm_delete.html"
    success_url = reverse_lazy('list-plp-record')

class PrintPlpRecord(View, LoginRequiredMixin):
    def get(self, request,pk):
        plp_rec = plprecord.objects.get(id=pk)
        plp = plps.objects.get(plp_code=plp_rec.plp_id)
        student = students.objects.get(nis_student=plp_rec.student_id)
        d_absensi ={}
        d_nilai =[]
        try:
            xl = pd.read_excel(plp_rec.report_excel,na_filter=False)
            d_absensi = {
                'alpha':str(int(xl['alpha'][0])),
                'sakit':str(int(xl['sakit'][0])),
                'izin':str(int(xl['izin'][0])),
            }
            nilai =  xl[['no','materi','submateri','teori','praktik']]
            # nilai = nilai.replace('', '-', regex=True)
            for item in nilai.to_numpy():
                d_nilai.append({
                    'no':str(item[0]),
                    'materi':item[1],
                    'sub_materi':item[2],
                    'teori':str(item[3]),
                    'praktik':str(item[4]),
                })
        except Exception as e:
            print(e)
        d_student = {
            'nis_student':student.nis_student,
            'name_student':student.name_student,
            'class_student':student.class_student,
            'sem_student':student.sem_student
        }
       
        document = MailMerge(plp.plp_rapor)
        # print(d_student)
        # print(d_nilai)
        # print(d_absensi)
        document.merge(**d_student)
        print(d_nilai)
        document.merge_rows('no', d_nilai)
        document.merge(**d_absensi)
        result = os.path.join(BASE_DIR, 'media/Rapor/'+student+"/"+plp.plp_name+"_"+str(student.class_student)+"_"+student.sem_student+".docx")
        a = document.write(result)
        print(type(a))
        plp_rec.report_result = result
        plp_rec.save()
        return redirect('list-plp-record')
#Test
class CheckAPi(View):
    def get(self, request):
        a = plps.objects.get(plp_code='P-01')
        path = a.plp_rapor
        print(path)
        document = MailMerge(path)
        print(document.get_merge_fields())
        # file_rapor = os.path.join(BASE_DIR, 'media',path)
        # print(file_rapor)
        return HttpResponse('OK')

#Ust Pur
class FillRapor(View):
    template = 'pages/form.html'
    def get(self, request, pk):
        plp_rec = plprecord.objects.get(id=pk)
        plp = plps.objects.get(plp_code=plp_rec.plp_id)
        student = students.objects.get(nis_student=plp_rec.student_id)
        document = None
        if student.sex_student=='L':
            document = MailMerge(plp.plp_rapor_qbs)
        else:
            document = MailMerge(plp.plp_rapor_fq)
        # document = MailMerge(plp.plp_rapor)
        f = list(document.get_merge_fields())
        f.remove('nis_student')
        f.remove('name_student')
        f.remove('sem_student')
        f.remove('class_student')
        f.remove('sakit')
        f.remove('izin')
        f.remove('alpha')
        f.sort()
        form = RaporForm(qs = f)
        return render(request, self.template,{'form':form,'button':'Submit','title':'Page Rapor '+student.name_student+'-'+plp.plp_name})
    
    def post(self, request, pk):
        body = request.POST.dict()
        body.pop('csrfmiddlewaretoken',None)
        # print(body)
        plp_rec = plprecord.objects.get(id=pk)
        plp = plps.objects.get(plp_code=plp_rec.plp_id)
        student = students.objects.get(nis_student=plp_rec.student_id)
        # document = MailMerge(plp.plp_rapor)
        document = None
        if student.sex_student=='L':
            document = MailMerge(plp.plp_rapor_qbs)
        else:
            document = MailMerge(plp.plp_rapor_fq)
        body['nis_student'] = student.nis_student
        body['name_student'] = student.name_student
        body['sem_student'] = student.sem_student
        body['class_student'] = student.class_student
        print(body)
        plp_rec.o_nilai = body
        plp_rec.save()
        # email_user = request.user.email
        # content = {
        #     'plp_id':plp.plp_code,
        #     'student_id': student.nis_student,
        #     'user_id':email_user,
        #     'o_nilai':body,

        # }
        # # o = content_rapor.objects.create(**content)
        # o.save()
        document.merge(**body)
        result = os.path.join(BASE_DIR, 'media/Students/'+str(student.nis_student)+"/"+plp.plp_name+"_"+str(student.class_student)+"_"+student.sem_student+".docx")
        a = document.write(result)
        print(type(a))
        plp_rec.report_result = result
        plp_rec.save()
        return redirect('list-plp-record')

#---------------------Subject----------------------------
class ListSubject(LoginRequiredMixin,SingleTableMixin, FilterView):
    login_url = '/login/'
    table_class = SubjectTable
    model = subjects
    queryset = subjects.objects.all()
    template_name = "pages/index.html"
    # filterset_class = PlpFilter
    def get_context_data(self, **kwargs):
        context = super(ListSubject, self).get_context_data(**kwargs)
        context['del'] = '/del-subject/'
        return context

class DeleteAllSubject(View):
    # models = parents
    # success_url = reverse_lazy('list-parent')
    # template_name_suffix = '_confirm_delete'
    # queryset = parents.objects.all()
    def post(self, request):
        _arr = request.POST.getlist('selection')
        for _id in _arr:
            subjects.objects.filter(id=_id).delete()
        return redirect('list-subject') 

class AddSubject(LoginRequiredMixin, CreateView):
    form_class = SubjectForm
    login_url = '/login/' 
    template_name = 'pages/form.html'
    success_url = reverse_lazy('list-subject')
    success_message = 'Doc successfully created!'
    error_message = "Error saving the Doc, check fields below."
    def get_context_data(self, **kwargs):
        context = super(AddSubject, self).get_context_data(**kwargs)
        context['title'] = 'Add Subject'
        context['button'] = 'Add Subject'
        return context

class UpdateSubject(UpdateView, LoginRequiredMixin):
    template_name = 'pages/form.html'
    success_url = reverse_lazy('list-subject')
    form_class = SubjectForm
    model = subjects
    login_url = '/login/'
    def get_context_data(self, **kwargs):
        context = super(UpdateSubject, self).get_context_data(**kwargs)
        context['title'] = 'Update Subject'
        context['button'] = 'Update Subject'
        return context
 
class DeleteSubject(LoginRequiredMixin, DeleteView):
    login_url = '/login/' 
    model = subjects
    template_name = "app/confirm_delete.html"
    success_url = reverse_lazy('list-subject')

#---------------------Subject Document----------------------------
class ListSubjectDoc(LoginRequiredMixin,SingleTableMixin, FilterView):
    login_url = '/login/'
    table_class = SubjectDocTable
    model = subject_doc
    queryset = subject_doc.objects.all()
    template_name = "pages/index.html"
    filterset_class = SubjectDocFilter
    def get_context_data(self, **kwargs):
        context = super(ListSubjectDoc, self).get_context_data(**kwargs)
        context['del'] = '/del-subject-doc/'
        return context

class DeleteAllSubjectDoc(View):
    # models = parents
    # success_url = reverse_lazy('list-parent')
    # template_name_suffix = '_confirm_delete'
    # queryset = parents.objects.all()
    def post(self, request):
        _arr = request.POST.getlist('selection')
        for _id in _arr:
            subject_doc.objects.filter(id=_id).delete()
        return redirect('list-subject-doc') 

class AddSubjectDoc(LoginRequiredMixin, CreateView):
    form_class = SubjectForm
    login_url = '/login/' 
    template_name = 'pages/form.html'
    success_url = reverse_lazy('list-subject-doc')
    success_message = 'Doc successfully created!'
    error_message = "Error saving the Doc, check fields below."
    def get_context_data(self, **kwargs):
        context = super(AddSubjectDoc, self).get_context_data(**kwargs)
        context['title'] = 'Add Subject Document'
        context['button'] = 'Add Subject Document'
        return context

class UpdateSubjectDoc(UpdateView, LoginRequiredMixin):
    template_name = 'pages/form.html'
    success_url = reverse_lazy('list-subject-doc')
    form_class = SubjectDocForm
    model = subject_doc
    login_url = '/login/'
    def get_context_data(self, **kwargs):
        context = super(UpdateSubjectDoc, self).get_context_data(**kwargs)
        context['title'] = 'Update Subject Document'
        context['button'] = 'Update Subject Document'
        return context
 
class DeleteSubjectDoc(LoginRequiredMixin, DeleteView):
    login_url = '/login/' 
    model = subject_doc
    template_name = "app/confirm_delete.html"
    success_url = reverse_lazy('list-subject-doc')

#-----------------------------------------EXPORT-----------------------------------------------
class ViewExportList(View):
    template = 'pages/export.html'
    data_plp = plps.objects.all()
    def get(self, request):     
        return render(request, self.template,{'plps':self.data_plp,'code':1,'link':0})
    def post(self, request):
        plp = request.POST['plp_code']
        s_date = request.POST['s_date']
        e_date = request.POST['e_date']
        rec = plprecord.objects.filter(plp_id=plp,join_plp__range=(s_date, e_date))
        data_json =[]
        for item in rec:
            a = {
                "nis":item.student_id,
                "date":item.join_plp
            }
            o = item.o_nilai.replace("\'", "\"")
            o = json.loads(o)
            # c = {key: value for (key, value) in (a.items() + json.loads(o).items())}
            c = {**a, **o}
            data_json.append(c)
            # print(type(a))
            # print(type(o))
        df = pd.DataFrame(data_json)
        a = df.to_excel(os.path.join(BASE_DIR, 'media/Export/'+str(plp)+"_"+str(s_date)+"_"+str(e_date)+".xls"))
        link = '/media/Export/'+str(plp)+"_"+str(s_date)+"_"+str(e_date)+".xls"
        return render(request, self.template,{'plps':self.data_plp,'code':1,'link':link})

class ViewExportBundle(View):
    template = 'pages/export.html'
    data_student = students.objects.all()
    def get(self, request):
        return render(request, self.template,{'students':self.data_student,'code':2,'link':0})
    def post(self, request):
        nis = request.POST['nis_student']
        s_date = request.POST['s_date']
        e_date = request.POST['e_date']
        rec = plprecord.objects.filter(student_id=nis,join_plp__range=(s_date, e_date)).order_by('position')
        docs = []
        link = '/media/Export/'+str(nis)+"_"+str(s_date)+"_"+str(e_date)+".docx"
        for item in rec:
            docs.append(str(item.report_result))
        merged_document = Document()
        print(docs)
        result = Document(docs[0])
        result.add_page_break()
        composer = Composer(result)
        for i in range(1, len(docs)):
            doc = Document(docs[i])
            if i != len(docs) - 1:
                doc.add_page_break()
            composer.append(doc)
        composer.save(os.path.join(BASE_DIR, 'media/Export/'+str(nis)+"_"+str(s_date)+"_"+str(e_date)+".docx"))
        # for index, file in enumerate(docs):
        #     sub_doc = Document(file)

        # # Don't add a page break if you've reached the last file.
        #     if index < len(docs)-1:
        #         sub_doc.add_page_break()

        #     for element in sub_doc.element.body:
        #         merged_document.element.body.append(element)
        # merged_document.save(os.path.join(BASE_DIR, 'media/Export/'+str(nis)+"_"+str(s_date)+"_"+str(e_date)+".docx"))
        
        return render(request, self.template,{'students':self.data_student,'code':2,'link':link})